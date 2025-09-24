import json
import re
import os
from typing import Dict, Any, List
from datetime import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def _ensure_url(url: str) -> str:
    if not url:
        return url
    u = url.strip()
    if u.startswith("mailto:") or u.startswith("http://") or u.startswith("https://"):
        return u
    if u.startswith("www."):
        return "https://" + u
    return "https://" + u


def _add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph."""
    url = _ensure_url(url)
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # underline + blue color
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    color = OxmlElement('w:color'); color.set(qn('w:val'), "0000FF"); rPr.append(color)

    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)
    return hyperlink


def _add_divider(doc):
    p = doc.add_paragraph()
    p_par = p._element
    pPr = p_par.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    p_borders.append(bottom)
    pPr.append(p_borders)
    return p


def _add_section_title(doc, title: str):
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def _safe_text(text: str) -> str:
    """Ensure text is safe for docx XML by removing any problematic characters."""
    if not text:
        return ""

    # Remove NULL bytes and control characters except for common whitespace
    sanitized = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)

    # Replace multiple whitespaces with single space
    sanitized = re.sub(r'\s+', ' ', sanitized)

    # Strip leading/trailing whitespace
    sanitized = sanitized.strip()

    return sanitized


def _safe_add_paragraph(doc, text: str):
    """Safely add a paragraph with text sanitization."""
    safe_text = _safe_text(text)
    if safe_text:
        return doc.add_paragraph(safe_text)
    return doc.add_paragraph("")


def render_harvard(resume_json, output_path: str, job_title: str = ""):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # tighten paragraph spacing globally for compact Harvard look
    for pstyle in doc.styles:
        try:
            pstyle.paragraph_format.space_after = Pt(2)
        except Exception:
            pass

    # --- Header: centered name + contact info ---
    name = resume_json.get("contact_info", {}).get("full_name", "")
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.add_run(name)
        r.bold = True
        r.font.size = Pt(16)
        p.paragraph_format.space_after = Pt(4)

    ci = resume_json.get("contact_info", {})
    links = resume_json.get("links", {}) or {}

    # Build contact line (exclude Coursera from header)
    cp = doc.add_paragraph()
    cp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_items = []

    if ci.get("location"): contact_items.append(ci["location"])
    if ci.get("email"): contact_items.append(ci["email"])
    if ci.get("phone"): contact_items.append(ci["phone"])

    # Add professional links to header (exclude Coursera)
    first = True
    if contact_items:
        cp.add_run(" • ".join(contact_items))
        first = False

    for label in ["LinkedIn", "GitHub", "HuggingFace"]:
        url = links.get(label)
        if url:
            if not first:
                cp.add_run(" • ")
            _add_hyperlink(cp, url, label)
            first = False

    cp.paragraph_format.space_after = Pt(6)
    _add_divider(doc)

    # === Summary ===
    if resume_json.get("summary"):
        _add_section_title(doc, "SUMMARY")
        doc.add_paragraph(resume_json["summary"])
        _add_divider(doc)

    # === Education ===
    if resume_json.get("education"):
        _add_section_title(doc, "EDUCATION")
        for edu in resume_json["education"]:
            table = doc.add_table(rows=1, cols=2)
            table.allow_autofit = True
            left_cell, right_cell = table.rows[0].cells

            # Left cell: Institution + location
            p = left_cell.paragraphs[0]
            run = p.add_run(edu.get("institution", ""))
            run.bold = True
            if edu.get("location"):
                p.add_run(f" — {edu['location']}")

            # Right cell: Graduation date
            if edu.get("graduation_date"):
                rp = right_cell.paragraphs[0]
                rp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                rp.add_run(edu["graduation_date"])

            # Next line: Degree + field + GPA
            degree_line = f"{edu.get('degree','')} in {edu.get('field','')}"
            if edu.get("gpa"):
                degree_line += f", GPA: {edu['gpa']}"
            left_cell.add_paragraph(degree_line)

        doc.add_paragraph()
        _add_divider(doc)

    # === Experience (table layout so dates align on right) ===
    if resume_json.get("experience"):
        _add_section_title(doc, "EXPERIENCE")
        for exp in resume_json["experience"]:
            table = doc.add_table(rows=1, cols=2)
            table.allow_autofit = True
            left_cell, right_cell = table.rows[0].cells

            # Left: company (bold) + position (italic) + bullets
            left_para = left_cell.paragraphs[0]
            left_run = left_para.add_run(_safe_text(exp.get("company", "")))
            left_run.bold = True

            # position on next line
            if exp.get("position"):
                pos_p = left_cell.add_paragraph()
                pos_run = pos_p.add_run(_safe_text(exp.get("position")))
                pos_run.italic = True

            # bullets
            for b in exp.get("achievements", []):
                bull_p = left_cell.add_paragraph()
                bull_p.add_run(f"• {_safe_text(b)}")
                bull_p.paragraph_format.space_after = Pt(1)

            # Right: dates / location (top-right)
            right_para = right_cell.paragraphs[0]
            date_text = ""
            if exp.get("start_date") or exp.get("end_date"):
                date_text = f"{exp.get('start_date','')} – {exp.get('end_date','')}"
            if exp.get("location"):
                date_text = (date_text + " | " if date_text else "") + exp.get("location")
            right_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            right_para.add_run(date_text)

        doc.add_paragraph()
        _add_divider(doc)

    # === Projects ===
    if resume_json.get("projects"):
        _add_section_title(doc, "PROJECTS")
        for proj in resume_json["projects"]:
            p = doc.add_paragraph()
            run = p.add_run(proj.get("title", ""))
            run.bold = True
            if proj.get("technologies"):
                p.add_run(f" — {', '.join(proj['technologies'])}")

            for b in proj.get("bullets", []):
                doc.add_paragraph(f"• {b}")
        _add_divider(doc)

    # === Certifications (including Coursera from links) ===
    certs = resume_json.get("certifications", []) or []

    # Add Coursera links from links dict to certifications section
    coursera_links = resume_json.get("links", {}).get("Coursera", [])
    if coursera_links:
        if isinstance(coursera_links, str):
            coursera_links = [coursera_links]
        certs.extend(coursera_links)

    if certs:
        _add_section_title(doc, "CERTIFICATIONS")
        for c in certs:
            doc.add_paragraph(str(c))
        doc.add_paragraph()
        _add_divider(doc)

    # === Skills & Languages ===
    skills = resume_json.get("skills", {})
    langs = resume_json.get("languages", [])

    if skills or langs:
        _add_section_title(doc, "SKILLS & INTERESTS")

        if skills:
            for cat, items in skills.items():
                doc.add_paragraph(f"{cat}: {', '.join(items)}")

        if langs:
            doc.add_paragraph("Languages: " + ", ".join(langs))

    # Save DOCX
    doc.save(output_path)


__all__ = ["render_harvard"]